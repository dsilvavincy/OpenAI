"""
Streamlit UI for AI-Driven T12 Data Question Generator
"""
import streamlit as st
import pandas as pd
import io
import os
import time
from datetime import datetime
from src.preprocess import tidy_sheet_all
from src.kpi_summary import generate_kpi_summary
from src.prompt import build_prompt, call_openai, validate_response
from src.output_quality import post_process_output
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

# Page configuration
st.set_page_config(
    page_title="T12 Property Analysis Tool",
    page_icon="🏢",
    layout="wide"
)

def validate_uploaded_file(uploaded_file):
    """Validate uploaded T12 file"""
    validation_results = {
        "is_valid": False,
        "messages": [],
        "warnings": []
    }
    
    # Check file extension
    if not uploaded_file.name.endswith(('.xlsx', '.xls')):
        validation_results["messages"].append("❌ Invalid file format. Please upload an Excel file (.xlsx or .xls)")
        return validation_results
    
    # Check file size (max 50MB)
    file_content = uploaded_file.getvalue()  # Get file content without moving file pointer
    file_size = len(file_content)
    
    if file_size > 50 * 1024 * 1024:  # 50MB
        validation_results["messages"].append("❌ File too large. Maximum size is 50MB")
        return validation_results
    
    if file_size < 1024:  # 1KB
        validation_results["messages"].append("❌ File too small. This doesn't appear to be a valid T12 file")
        return validation_results
    
    validation_results["is_valid"] = True
    validation_results["messages"].append("✅ File format and size validation passed")
    
    if file_size > 10 * 1024 * 1024:  # 10MB
        validation_results["warnings"].append("⚠️ Large file detected. Processing may take longer")
    
    return validation_results

def create_progress_tracker():
    """Create a progress tracking system"""
    if 'processing_steps' not in st.session_state:
        st.session_state.processing_steps = {
            'file_upload': False,
            'file_validation': False, 
            'data_processing': False,
            'kpi_generation': False,
            'ai_analysis': False
        }
    return st.session_state.processing_steps

def update_progress(step_name, status=True):
    """Update progress for a specific step"""
    if 'processing_steps' in st.session_state:
        st.session_state.processing_steps[step_name] = status

def display_progress():
    """Display progress indicators"""
    progress = st.session_state.get('processing_steps', {})
    
    st.subheader("🔄 Processing Progress")
    
    steps = [
        ('file_upload', '📁 File Upload'),
        ('file_validation', '✅ File Validation'),
        ('data_processing', '⚙️ Data Processing'),
        ('kpi_generation', '📊 KPI Generation'),
        ('ai_analysis', '🤖 AI Analysis')
    ]
    
    for step_key, step_label in steps:
        status = progress.get(step_key, False)
        if status:
            st.success(f"{step_label} ✓")
        else:
            st.info(f"{step_label} ⏳")

def generate_pdf_report(processed_output):
    """Generate a professional PDF report from processed output"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=20,
        textColor=colors.darkblue
    )
    
    # Build content
    story = []
    
    # Title
    story.append(Paragraph("T12 PROPERTY ANALYSIS REPORT", title_style))
    story.append(Spacer(1, 20))
    
    # Property Information
    property_info = processed_output["property_info"]
    story.append(Paragraph("PROPERTY INFORMATION", heading_style))
    property_data = [
        ["Property Name:", property_info.get('name', 'N/A')],
        ["Property Address:", property_info.get('address', 'N/A')],
        ["Report Generated:", datetime.fromisoformat(processed_output["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")]
    ]
    property_table = Table(property_data, colWidths=[2*inch, 4*inch])
    property_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
    ]))
    story.append(property_table)
    story.append(Spacer(1, 20))
    
    # Analysis Summary
    quality = processed_output["quality_metrics"]
    analysis = processed_output["analysis"]
    
    story.append(Paragraph("ANALYSIS SUMMARY", heading_style))
    summary_data = [
        ["Quality Score:", f"{quality['overall_score']}/100 ({quality['quality_level']})"],
        ["Strategic Questions:", str(len(analysis['strategic_questions']))],
        ["Recommendations:", str(len(analysis['recommendations']))],
        ["Concerning Trends:", str(len(analysis['concerning_trends']))]
    ]
    summary_table = Table(summary_data, colWidths=[2*inch, 4*inch])
    summary_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 20))
    
    # Strategic Questions
    story.append(Paragraph("STRATEGIC MANAGEMENT QUESTIONS", heading_style))
    for i, question in enumerate(analysis["strategic_questions"], 1):
        story.append(Paragraph(f"{i}. {question}", styles['Normal']))
        story.append(Spacer(1, 10))
    
    story.append(Spacer(1, 20))
    
    # Recommendations
    story.append(Paragraph("ACTIONABLE RECOMMENDATIONS", heading_style))
    for i, rec in enumerate(analysis["recommendations"], 1):
        story.append(Paragraph(f"{i}. {rec}", styles['Normal']))
        story.append(Spacer(1, 10))
    
    story.append(Spacer(1, 20))
    
    # Concerning Trends
    story.append(Paragraph("CONCERNING TRENDS", heading_style))
    for i, concern in enumerate(analysis["concerning_trends"], 1):
        story.append(Paragraph(f"{i}. {concern}", styles['Normal']))
        story.append(Spacer(1, 10))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_word_report(processed_output):
    """Generate a professional Word document report from processed output"""
    doc = Document()
    
    # Title
    title = doc.add_heading('T12 PROPERTY ANALYSIS REPORT', 0)
    title.alignment = 1  # Center alignment
    
    # Property Information
    doc.add_heading('PROPERTY INFORMATION', level=1)
    property_info = processed_output["property_info"]
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells
    cells[0].text = 'Property Name:'
    cells[1].text = property_info.get('name', 'N/A')
    
    cells = table.rows[1].cells
    cells[0].text = 'Property Address:'
    cells[1].text = property_info.get('address', 'N/A')
    
    cells = table.rows[2].cells
    cells[0].text = 'Report Generated:'
    cells[1].text = datetime.fromisoformat(processed_output["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")
    
    doc.add_paragraph()
    
    # Analysis Summary
    doc.add_heading('ANALYSIS SUMMARY', level=1)
    quality = processed_output["quality_metrics"]
    analysis = processed_output["analysis"]
    
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_data = [
        ['Quality Score:', f"{quality['overall_score']}/100 ({quality['quality_level']})"],
        ['Strategic Questions:', str(len(analysis['strategic_questions']))],
        ['Recommendations:', str(len(analysis['recommendations']))],
        ['Concerning Trends:', str(len(analysis['concerning_trends']))]
    ]
    
    for i, (key, value) in enumerate(summary_data):
        cells = summary_table.rows[i].cells
        cells[0].text = key
        cells[1].text = value
    
    doc.add_paragraph()
    
    # Strategic Questions
    doc.add_heading('STRATEGIC MANAGEMENT QUESTIONS', level=1)
    for i, question in enumerate(analysis["strategic_questions"], 1):
        doc.add_paragraph(f"{i}. {question}")
    
    # Recommendations
    doc.add_heading('ACTIONABLE RECOMMENDATIONS', level=1)
    for i, rec in enumerate(analysis["recommendations"], 1):
        doc.add_paragraph(f"{i}. {rec}")
    
    # Concerning Trends
    doc.add_heading('CONCERNING TRENDS', level=1)
    for i, concern in enumerate(analysis["concerning_trends"], 1):
        doc.add_paragraph(f"{i}. {concern}")
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def main():
    st.title("🏢 AI-Driven T12 Property Analysis Tool")
    st.markdown("Upload your T12 Excel file to generate AI-powered property performance insights")
    
    # Initialize progress tracking
    create_progress_tracker()
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # OpenAI API Key input with validation
        api_key = st.text_input(
            "OpenAI API Key", 
            type="password",
            help="Enter your OpenAI API key to enable AI analysis"
        )
        
        if api_key:
            if api_key.startswith('sk-') and len(api_key) > 20:
                st.success("✅ API key format looks valid")
            else:
                st.warning("⚠️ API key format may be incorrect")
        
        st.divider()
        
        # Property information
        st.header("🏢 Property Information")
        property_name = st.text_input("Property Name", placeholder="e.g., Sunset Apartments")
        property_address = st.text_input("Property Address", placeholder="e.g., 123 Main St, City, State")
        
        st.divider()
        
        # Display progress
        display_progress()
        
    # Main content area
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.header("📁 Upload T12 Data")
        
        uploaded_file = st.file_uploader(
            "Choose Excel file",
            type=['xlsx', 'xls'],
            help="Upload your T12 property financial data in Excel format"
        )
        
        if uploaded_file is not None:
            update_progress('file_upload', True)
            
            # File validation
            validation = validate_uploaded_file(uploaded_file)
            
            # Display validation results
            for msg in validation["messages"]:
                if msg.startswith("✅"):
                    st.success(msg)
                elif msg.startswith("❌"):
                    st.error(msg)
                    
            for warning in validation["warnings"]:
                st.warning(warning)
            
            if not validation["is_valid"]:
                return
                
            update_progress('file_validation', True)
            
            try:
                # Show file details
                file_size = len(uploaded_file.getvalue())
                st.info(f"� **File:** {uploaded_file.name}")
                st.info(f"📏 **Size:** {file_size/1024:.1f} KB")
                
                # Process the file with progress indicator
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                status_text.text("📂 Reading Excel file...")
                progress_bar.progress(20)
                
                # Read Excel file directly from memory (no temp file)
                status_text.text("⚙️ Processing T12 data...")
                progress_bar.progress(50)
                
                # Use BytesIO to pass file-like object to tidy_sheet_all
                excel_buffer = io.BytesIO(uploaded_file.getvalue())
                df = tidy_sheet_all(excel_buffer)
                
                status_text.text("✅ Data processing complete!")
                progress_bar.progress(100)
                
                update_progress('data_processing', True)
                
                st.success(f"✅ Successfully processed {len(df):,} data points")
                
                # Show data preview
                with st.expander("📊 View Data Preview", expanded=False):
                    st.dataframe(df.head(10), use_container_width=True)
                    st.info(f"Showing first 10 rows of {len(df):,} total rows")

                # Show rows with invalid MonthParsed for debugging
                if 'MonthParsed' in df.columns and 'IsYTD' in df.columns:
                    # Only check non-YTD rows for invalid MonthParsed
                    non_ytd_df = df[df['IsYTD'] == False]
                    invalid_month_rows = non_ytd_df[non_ytd_df['MonthParsed'].isnull()]
                    
                    with st.expander(f"🚨 Invalid MonthParsed (Non-YTD): {len(invalid_month_rows)}", expanded=False):
                        if len(invalid_month_rows) > 0:
                            st.dataframe(invalid_month_rows, use_container_width=True)
                            st.error(f"❌ {len(invalid_month_rows)} non-YTD rows have invalid MonthParsed")
                            
                            # Show unique Month values that failed to parse
                            unique_failed_months = invalid_month_rows['Month'].unique()
                            st.write("**Failed to parse these Month values:**")
                            for month in unique_failed_months[:10]:  # Show first 10
                                st.write(f"- '{month}'")
                            if len(unique_failed_months) > 10:
                                st.write(f"... and {len(unique_failed_months) - 10} more")
                        else:
                            st.success("✅ All non-YTD MonthParsed values are valid!")
                    
                    # Show YTD summary separately
                    ytd_df = df[df['IsYTD'] == True]
                    with st.expander(f"📅 YTD Rows Summary: {len(ytd_df)}", expanded=False):
                        if len(ytd_df) > 0:
                            st.info(f"✅ {len(ytd_df)} YTD rows found (MonthParsed is expected to be null)")
                            st.dataframe(ytd_df[['Metric', 'Month', 'Value']].head(10), use_container_width=True)
                        else:
                            st.warning("⚠️ No YTD rows found in data")
                
                # Show DataFrame structure and statistics
                with st.expander("📊 DataFrame Analysis", expanded=False):
                    col_info1, col_info2 = st.columns(2)
                    
                    with col_info1:
                        st.write("**Shape:**", df.shape)
                        st.write("**Columns:**", list(df.columns))
                        st.write("**Data Types:**")
                        st.dataframe(pd.DataFrame(df.dtypes).rename(columns={0: 'Type'}), use_container_width=True)
                    
                    with col_info2:
                        st.write("**Key Statistics:**")
                        if 'IsYTD' in df.columns:
                            ytd_count = df[df['IsYTD'] == True].shape[0]
                            non_ytd_count = df[df['IsYTD'] == False].shape[0]
                            st.write(f"- YTD rows: {ytd_count}")
                            st.write(f"- Monthly rows: {non_ytd_count}")
                        
                        if 'Metric' in df.columns:
                            st.write(f"- Unique metrics: {df['Metric'].nunique()}")
                            st.write("**Top 10 Metrics:**")
                            top_metrics = df['Metric'].value_counts().head(10)
                            for metric, count in top_metrics.items():
                                st.write(f"  • {metric}: {count} rows")
                        
                        if 'Value' in df.columns:
                            st.write(f"- Total values: {df['Value'].count()}")
                            st.write(f"- Missing values: {df['Value'].isna().sum()}")
                
                # Test KPI Summary Generation
                with st.expander("🧪 Test KPI Summary Generation", expanded=False):
                    if st.button("Generate Test KPI Summary", key="test_kpi"):
                        try:
                            test_kpi_summary = generate_kpi_summary(df)
                            st.text_area("Generated KPI Summary:", test_kpi_summary, height=400)
                            st.success("✅ KPI Summary generated successfully!")
                        except Exception as e:
                            st.error(f"❌ Error generating KPI Summary: {str(e)}")
                            st.write("**DataFrame info for debugging:**")
                            st.write(f"Shape: {df.shape}")
                            st.write(f"Columns: {list(df.columns)}")
                            if not df.empty:
                                st.write("**Sample data:**")
                                st.dataframe(df.head(), use_container_width=True)
                
            except Exception as e:
                st.error(f"❌ Error processing file: {str(e)}")
                st.info("💡 **Troubleshooting tips:**")
                st.info("• Ensure the file is a valid T12 Excel format")
                st.info("• Check that the file isn't corrupted")
                st.info("• Try uploading a different file")
                return
    
    with col2:
        st.header("📈 Analysis Results")
        
        if uploaded_file is not None and 'df' in locals():
            try:
                # Generate KPI Summary with progress
                kpi_progress = st.progress(0)
                kpi_status = st.empty()
                
                kpi_status.text("📊 Generating KPI summary...")
                kpi_progress.progress(50)
                
                kpi_summary = generate_kpi_summary(df)
                
                kpi_status.text("✅ KPI summary complete!")
                kpi_progress.progress(100)
                
                update_progress('kpi_generation', True)
                
                # Show KPI Summary
                st.subheader("📋 KPI Summary")
                with st.expander("View KPI Summary", expanded=True):
                    st.text_area("", kpi_summary, height=300, disabled=True)
                
                # Test OpenAI Prompt Generation
                with st.expander("🧪 Test OpenAI Prompt", expanded=False):
                    if st.button("Generate Test Prompt", key="test_prompt"):
                        try:
                            system_prompt, user_prompt = build_prompt(kpi_summary)
                            st.write("**System Prompt:**")
                            st.text_area("", system_prompt, height=200, disabled=True, key="sys_prompt")
                            st.write("**User Prompt:**")
                            st.text_area("", user_prompt, height=300, disabled=True, key="user_prompt")
                            st.success("✅ Prompts generated successfully!")
                            
                            # Show token count estimation
                            total_chars = len(system_prompt) + len(user_prompt)
                            estimated_tokens = total_chars // 4  # Rough estimation
                            st.info(f"📊 Estimated tokens: ~{estimated_tokens:,} (Character count: {total_chars:,})")
                            
                        except Exception as e:
                            st.error(f"❌ Error generating prompts: {str(e)}")
                
                # AI Analysis Section
                st.subheader("🤖 AI-Powered Analysis")
                
                if not api_key:
                    st.warning("⚠️ Please enter your OpenAI API key in the sidebar to generate AI analysis")
                    st.info("💡 Your API key is used securely and is not stored")
                
                elif st.button("� Generate AI Analysis", type="primary", use_container_width=True):
                    # AI Analysis with detailed progress
                    ai_progress = st.progress(0)
                    ai_status = st.empty()
                    
                    ai_status.text("🔧 Building analysis prompt...")
                    ai_progress.progress(25)
                    
                    system_prompt, user_prompt = build_prompt(kpi_summary)
                    
                    ai_status.text("🤖 Calling OpenAI API...")
                    ai_progress.progress(50)
                    
                    ai_response = call_openai(system_prompt, user_prompt, api_key)
                    
                    ai_status.text("✨ Processing AI response...")
                    ai_progress.progress(75)
                    
                    # Validate and post-process response
                    if not ai_response.startswith("Error:"):
                        is_valid, validation_msg = validate_response(ai_response)
                        
                        if is_valid:
                            # Post-process the output
                            property_info = {
                                "name": property_name or "Unknown Property",
                                "address": property_address or "No address provided"
                            }
                            
                            processed_output = post_process_output(ai_response, property_info)
                            
                            ai_status.text("✅ Analysis complete!")
                            ai_progress.progress(100)
                            
                            update_progress('ai_analysis', True)
                            
                            # Display formatted results
                            st.markdown("### 📊 Analysis Results")
                            st.markdown(processed_output["display_text"])
                            
                            # Show detailed AI analysis sections
                            analysis = processed_output["analysis"]
                            
                            # Strategic Questions Section
                            with st.expander("🎯 Strategic Management Questions", expanded=True):
                                for i, question in enumerate(analysis["strategic_questions"], 1):
                                    st.markdown(f"**{i}.** {question}")
                            
                            # Recommendations Section
                            with st.expander("💡 Actionable Recommendations", expanded=True):
                                for i, rec in enumerate(analysis["recommendations"], 1):
                                    st.markdown(f"**{i}.** {rec}")
                            
                            # Concerning Trends Section
                            with st.expander("⚠️ Concerning Trends", expanded=True):
                                for i, concern in enumerate(analysis["concerning_trends"], 1):
                                    st.markdown(f"**{i}.** {concern}")
                            
                            # Show quality metrics
                            quality = processed_output["quality_metrics"]
                            col_q1, col_q2, col_q3 = st.columns(3)
                            
                            with col_q1:
                                st.metric("Quality Score", f"{quality['overall_score']}/100")
                            with col_q2:
                                st.metric("Quality Level", quality['quality_level'])
                            with col_q3:
                                st.metric("Questions Found", len(processed_output["analysis"]["strategic_questions"]))
                            
                            # Export options
                            st.subheader("📄 Export Options")
                            
                            # Generate enhanced report content
                            report_content = generate_enhanced_report(processed_output)
                            pdf_content = generate_pdf_report(processed_output)
                            word_content = generate_word_report(processed_output)
                            
                            # Create filename base
                            filename_base = f"T12_Analysis_{property_name.replace(' ', '_') if property_name else 'Property'}_{datetime.now().strftime('%Y%m%d_%H%M')}"
                            
                            col_export1, col_export2, col_export3, col_export4 = st.columns(4)
                            
                            with col_export1:
                                st.download_button(
                                    label="📄 Download PDF Report",
                                    data=pdf_content,
                                    file_name=f"{filename_base}.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                            
                            with col_export2:
                                st.download_button(
                                    label="📝 Download Word Report",
                                    data=word_content,
                                    file_name=f"{filename_base}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                            
                            with col_export3:
                                st.download_button(
                                    label="📄 Download Text Report",
                                    data=report_content,
                                    file_name=f"{filename_base}.txt",
                                    mime="text/plain",
                                    use_container_width=True
                                )
                            
                            with col_export4:
                                st.download_button(
                                    label="� Download JSON Data",
                                    data=str(processed_output),
                                    file_name=f"{filename_base}.json",
                                    mime="application/json",
                                    use_container_width=True
                                )
                        else:
                            st.error(f"❌ Response validation failed: {validation_msg}")
                            st.warning("🔄 Try regenerating the analysis or check your API key")
                    else:
                        st.error(f"❌ {ai_response}")
                        if "api key" in ai_response.lower():
                            st.info("💡 Please check your OpenAI API key in the sidebar")
                        elif "rate limit" in ai_response.lower():
                            st.info("💡 Please wait a moment and try again")
                    
            except Exception as e:
                st.error(f"❌ Error during analysis: {str(e)}")
                st.info("💡 **Troubleshooting tips:**")
                st.info("• Check that your data was processed correctly")
                st.info("• Verify your OpenAI API key is valid")
                st.info("• Try refreshing the page and uploading again")
        else:
            st.info("👆 Please upload a T12 Excel file to begin analysis")

def generate_enhanced_report(processed_output):
    """Generate an enhanced formatted report from processed output"""
    analysis = processed_output["analysis"]
    metadata = processed_output["metadata"]
    property_info = processed_output["property_info"]
    quality = processed_output["quality_metrics"]
    
    timestamp = datetime.fromisoformat(processed_output["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")
    
    report = f"""
T12 PROPERTY ANALYSIS REPORT
============================

PROPERTY INFORMATION
====================
Property Name: {property_info.get('name', 'N/A')}
Property Address: {property_info.get('address', 'N/A')}
Report Generated: {timestamp}

ANALYSIS SUMMARY
================
Quality Score: {quality['overall_score']}/100 ({quality['quality_level']})
Strategic Questions: {len(analysis['strategic_questions'])}
Recommendations: {len(analysis['recommendations'])}
Concerning Trends: {len(analysis['concerning_trends'])}

STRATEGIC MANAGEMENT QUESTIONS
==============================
"""
    
    for i, question in enumerate(analysis["strategic_questions"], 1):
        report += f"{i}. {question}\n"
    
    report += f"""
ACTIONABLE RECOMMENDATIONS
==========================
"""
    
    for i, rec in enumerate(analysis["recommendations"], 1):
        report += f"{i}. {rec}\n"
    
    report += f"""
CONCERNING TRENDS
=================
"""
    
    for i, concern in enumerate(analysis["concerning_trends"], 1):
        report += f"{i}. {concern}\n"
    
    if quality["recommendations"]:
        report += f"""
QUALITY IMPROVEMENT SUGGESTIONS
===============================
"""
        for rec in quality["recommendations"]:
            report += f"• {rec}\n"
    
    report += f"""

TECHNICAL DETAILS
=================
Response Length: {metadata['response_length']} characters
Processing Timestamp: {metadata['processing_timestamp']}

---
Generated by AI-Driven T12 Analysis Tool
Quality Control System v1.0
"""
    
    return report

def generate_report(property_name, property_address, kpi_summary, ai_analysis):
    """Generate a simple formatted report for download (legacy function)"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    report = f"""
T12 PROPERTY ANALYSIS REPORT
============================

Property: {property_name}
Address: {property_address}
Generated: {timestamp}

KPI SUMMARY
===========
{kpi_summary}

AI ANALYSIS
===========
{ai_analysis}

---
Generated by AI-Driven T12 Analysis Tool
"""
    return report

if __name__ == "__main__":
    main()
