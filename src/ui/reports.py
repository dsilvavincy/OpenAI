"""
Report generation utilities for PDF, Word, and text exports
"""
import io
from datetime import datetime
from reportlab.lib.pagesizes import landscape, letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_ORIENT
import pandas as pd
from src.core.report_generator import ReportGenerator

# --- COLORS ---
# Hex to RGB Tuple for ReportLab
def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    return tuple(int(hex_str[i:i+2], 16)/255.0 for i in (0, 2, 4))

# --- COLORS (DARK MODE MATCHING CSS) ---
# Hex to RGB Tuple for ReportLab
def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    return tuple(int(hex_str[i:i+2], 16)/255.0 for i in (0, 2, 4))

COLOR_HEADER_TEAL = hex_to_rgb("#009879") 
COLOR_BODY_BG = hex_to_rgb("#262730") # Dark Body
COLOR_BODY_TEXT = hex_to_rgb("#ffffff") # White Text

# High Contrast Status Colors (Flip from previous: Dark BG, Light Text)
COLOR_TEXT_GREEN = hex_to_rgb("#a3e6b1")
COLOR_BG_GREEN = hex_to_rgb("#1e4620")
COLOR_TEXT_RED = hex_to_rgb("#f8d7da") 
COLOR_BG_RED = hex_to_rgb("#5a1a1e")
COLOR_TEXT_YELLOW = hex_to_rgb("#fdf2ce")
COLOR_BG_YELLOW = hex_to_rgb("#4d3e04")

COLOR_BORDER = hex_to_rgb("#444444")

def generate_pdf_report(processed_output, visual_data=None):
    """Generate a professional PDF report from processed output + visual tables"""
    buffer = io.BytesIO()
    # Use Landscape Letter
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), rightMargin=36, leftMargin=36, topMargin=50, bottomMargin=50)
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=20,
        alignment=1
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=12,
        textColor=colors.darkblue
    )
    metric_style = ParagraphStyle('Metric', parent=styles['Normal'], fontSize=8)
    
    story = []
    
    # Title
    story.append(Paragraph("T12 PROPERTY ANALYSIS REPORT", title_style))
    story.append(Spacer(1, 10))
    
    # Property Info
    property_info = processed_output.get("property_info", {})
    story.append(Paragraph("PROPERTY INFORMATION", heading_style))
    p_data = [
        ["Property Name:", property_info.get('name', 'N/A')],
        ["Report Period:", property_info.get('report_period', 'N/A')],
        ["Report Generated:", datetime.now().strftime("%Y-%m-%d %H:%M")]
    ]
    p_table = Table(p_data, colWidths=[2*inch, 4*inch])
    p_table.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
    ]))
    story.append(p_table)
    story.append(Spacer(1, 20))
    
    # --- VISUAL SECTIONS ---
    if visual_data:
        # 0. Portfolio Snapshot
        if 'portfolio_snapshot' in visual_data:
            p_snap = visual_data['portfolio_snapshot']
            if hasattr(p_snap, 'columns'):
                # Split logic matching UI
                cols = []
                for c in p_snap.columns:
                     c_norm = str(c).replace('\n', ' ').replace('  ', ' ').strip()
                     if "In Place Eff. Rate Prior Month" not in c_norm:
                         cols.append(c)
                
                # Slices (Indices based on filtered columns)
                # Group 1: 0-5 (Details)
                # Group 2: 5-10 (Ops)
                # Group 3: 10-16 (NOI)
                # Group 4: 16-22 (Rev)
                # Group 5: 22+ (Exp)
                
                story.append(Spacer(1, 20))
                
                # Slices (Indices based on filtered columns)
                # Group 1: 0-5 (Details)
                # Group 2: 5-10 (Ops)
                # Group 3: 10-16 (NOI)
                # Group 4: 16-22 (Rev)
                # Group 5: 22+ (Exp)
                
                groups = [
                    (cols[:5], "Property Details", COLOR_HEADER_TEAL),
                    (cols[5:10], "Cur. Mnth. Operations - Financial Based", COLOR_HEADER_TEAL),
                    (cols[10:16], "NOI - % Variance", COLOR_HEADER_TEAL),
                    (cols[16:22], "Revenue - % Variance", COLOR_HEADER_TEAL),
                    (cols[22:], "Expenses - % Variance", COLOR_HEADER_TEAL)
                ]
                
                for g_cols, title, hdr_bg in groups:
                    if not g_cols: continue
                    
                    story.append(Paragraph(title.upper(), heading_style))
                    # Extract subset
                    sub_df = p_snap[g_cols]
                    data = [sub_df.columns.tolist()] + sub_df.values.tolist()
                    
                    t = Table(data)
                    
                    # Base styles (DARK MODE)
                    base_style = [
                        ('FONTSIZE', (0,0), (-1,-1), 8),
                        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                        ('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER),
                        ('BACKGROUND', (0,0), (-1,-1), COLOR_BODY_BG), # Entire Table Dark
                        ('BACKGROUND', (0,0), (-1,0), hdr_bg), # Header Override
                        ('TEXTCOLOR', (0,0), (-1,-1), COLOR_BODY_TEXT), # White Text Default
                        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                    ]
                    
                    # Conditional Formatting Loop
                    # Skip header row (idx 0), so data row i corresponds to table row i+1
                    for r_idx, row in enumerate(sub_df.itertuples(index=False)):
                        for c_idx, val in enumerate(row):
                            col_name = str(sub_df.columns[c_idx])
                            
                            # Logic copied from ReportGenerator
                            try:
                                raw_val = float(val) if isinstance(val, (int, float)) else 0
                                is_num = isinstance(val, (int, float))
                            except: 
                                is_num = False
                                
                            if is_num:
                                text_col = COLOR_BODY_TEXT
                                bg_col = None # Default
                                
                                # Green/Red Logic (Matches HTML CSS Rules)
                                if "Physical Occupancy" in col_name:
                                    bg_col = COLOR_BG_GREEN if raw_val >= 0.90 else COLOR_BG_YELLOW if raw_val >= 0.85 else COLOR_BG_RED
                                    text_col = COLOR_TEXT_GREEN if raw_val >= 0.90 else COLOR_TEXT_YELLOW if raw_val >= 0.85 else COLOR_TEXT_RED
                                elif "Economic Occupancy" in col_name:
                                    bg_col = COLOR_BG_GREEN if raw_val >= 0.85 else COLOR_BG_YELLOW if raw_val >= 0.75 else COLOR_BG_RED
                                    text_col = COLOR_TEXT_GREEN if raw_val >= 0.85 else COLOR_TEXT_YELLOW if raw_val >= 0.75 else COLOR_TEXT_RED
                                elif "Debt Yield" in col_name:
                                    # Heuristic check
                                    cut_g, cut_y = (7.5, 5.95) if raw_val > 1 else (0.075, 0.0595)
                                    bg_col = COLOR_BG_GREEN if raw_val >= cut_g else COLOR_BG_YELLOW if raw_val >= cut_y else COLOR_BG_RED
                                    text_col = COLOR_TEXT_GREEN if raw_val >= cut_g else COLOR_TEXT_YELLOW if raw_val >= cut_y else COLOR_TEXT_RED
                                elif "DSCR" in col_name:
                                    bg_col = COLOR_BG_GREEN if raw_val >= 1.15 else COLOR_BG_YELLOW if raw_val >= 1.0 else COLOR_BG_RED
                                    text_col = COLOR_TEXT_GREEN if raw_val >= 1.15 else COLOR_TEXT_YELLOW if raw_val >= 1.0 else COLOR_TEXT_RED
                                elif "vs Bdgt" in col_name:
                                    cut_g, cut_y = (3.0, 0.0) if raw_val > 2 else (0.03, 0.0)
                                    bg_col = COLOR_BG_GREEN if raw_val >= cut_g else COLOR_BG_YELLOW if raw_val >= cut_y else COLOR_BG_RED
                                    text_col = COLOR_TEXT_GREEN if raw_val >= cut_g else COLOR_TEXT_YELLOW if raw_val >= cut_y else COLOR_TEXT_RED
                                elif "vs T1 Prior" in col_name or "Sequential" in col_name:
                                     # Variances
                                     if raw_val >= 0.01: text_col = COLOR_TEXT_GREEN # Simple text color for these? No HTML uses formatting too? 
                                     # HTML uses arrows and color spans. Let's use simple colors for PDF text.
                                     elif raw_val <= -0.01: text_col = COLOR_TEXT_RED
                                     else: text_col = colors.silver # Grayish
                                     
                                if bg_col:
                                    base_style.append(('BACKGROUND', (c_idx, r_idx+1), (c_idx, r_idx+1), bg_col))
                                    base_style.append(('TEXTCOLOR', (c_idx, r_idx+1), (c_idx, r_idx+1), text_col))
                                elif text_col != COLOR_BODY_TEXT:
                                    base_style.append(('TEXTCOLOR', (c_idx, r_idx+1), (c_idx, r_idx+1), text_col))

                    t.setStyle(TableStyle(base_style))
                    story.append(t)
                    story.append(Spacer(1, 10))
                
                story.append(Spacer(1, 20))
                
        # 1. KPI Snapshot
        if 'kpi_data' in visual_data:
            story.append(Paragraph("KPI SNAPSHOT", heading_style))
            kpi_t = visual_data['kpi_data']
            # Convert dict/df to list of lists for ReportLab
            if hasattr(kpi_t, 'columns'): # It's a dataframe
                data = [kpi_t.columns.tolist()] + kpi_t.values.tolist()
            else:
                data = kpi_t # Assuming it's already list of lists
                
            # Basic table style
            base_style = [
                ('FONTSIZE', (0,0), (-1,-1), 8),
                ('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER),
                ('BACKGROUND', (0,0), (-1,-1), COLOR_BODY_BG),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('BACKGROUND', (0,0), (-1,0), COLOR_HEADER_TEAL), # Dark Grey Header for KPI
                ('TEXTCOLOR', (0,0), (-1,-1), COLOR_BODY_TEXT),
            ]
            t.setStyle(TableStyle(base_style))
            story.append(t)
            story.append(Spacer(1, 20))

        # 2. Financial Data
        if 'financial_data' in visual_data:
            story.append(Paragraph("MONTHLY FINANCIAL DATA", heading_style))
            fin_t = visual_data['financial_data']
            # Convert dataframe to list of lists
            if hasattr(fin_t, 'columns'):
                # Wrap headers
                headers = [Paragraph(str(c), metric_style) for c in fin_t.columns]
                rows = []
                for row_idx, row in fin_t.iterrows():
                    r_data = []
                    for item in row:
                        if isinstance(item, (int, float)):
                            r_data.append(f"{item:,.0f}")
                        else:
                            r_data.append(Paragraph(str(item), metric_style))
                    rows.append(r_data)
                
                data = [headers] + rows
                
                # Auto-width calculation (simplistic)
                col_w = [0.8*inch] + [0.5*inch]*(len(headers)-1)
                
                t = Table(data, colWidths=col_w, repeatRows=1)
                t.setStyle(TableStyle([
                    ('FONTSIZE', (0,0), (-1,-1), 7), # Smaller for financial fits
                    ('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER),
                    ('BACKGROUND', (0,0), (-1,-1), COLOR_BODY_BG),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('BACKGROUND', (0,0), (-1,0), COLOR_HEADER_TEAL), # Dark Grey
                    ('TEXTCOLOR', (0,0), (-1,-1), COLOR_BODY_TEXT),
                ]))
                story.append(t)
                story.append(Spacer(1, 20))

    # --- AI ANALYSIS ---
    story.append(Paragraph("AI ANALYSIS & RECOMMENDATIONS", heading_style))
    
    # Budget Variances
    if 'budget_variances' in processed_output:
        story.append(Paragraph("Budget Variances", heading_style))
        bv = processed_output['budget_variances']
        for cat, items in bv.items():
            if not items: continue
            story.append(Paragraph(f"{cat}", styles['Heading3']))
            for item in items:
                text = f"<b>{item.get('metric')}</b><br/>"
                text += f"Actual: ${item.get('actual', 0):,} | Budget: ${item.get('budget', 0):,} | Var: {item.get('variance_pct', 0)}%<br/>"
                for q in item.get('questions', []):
                    text += f"• {q}<br/>"
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 5))
            story.append(Spacer(1, 10))
            
    # Trailing Anomalies
    if 'trailing_anomalies' in processed_output:
        story.append(Paragraph("Trailing Anomalies", heading_style))
        ta = processed_output['trailing_anomalies']
        for cat, items in ta.items():
            if not items: continue
            story.append(Paragraph(f"{cat}", styles['Heading3']))
            for item in items:
                text = f"<b>{item.get('metric')}</b><br/>"
                text += f"Current: ${item.get('current', 0):,} | T3 Avg: ${item.get('t3_avg', 0):,} | Dev: {item.get('deviation_pct', 0)}%<br/>"
                for q in item.get('questions', []):
                    text += f"• {q}<br/>"
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 5))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_word_report(processed_output, visual_data=None):
    """Generate a professional Word document report from processed output + visual tables"""
    doc = Document()
    
    # Title
    # TITLE & SETUP
    # LANDSCAPE ORIENTATION
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    title = doc.add_heading('T12 PROPERTY ANALYSIS REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Property Info
    doc.add_heading('PROPERTY INFORMATION', level=1)
    property_info = processed_output.get("property_info", {})
    p = doc.add_paragraph()
    p.add_run(f"Property Name: {property_info.get('name', 'N/A')}\n").bold = True
    p.add_run(f"Report Period: {property_info.get('report_period', 'N/A')}\n")
    p.add_run(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # --- VISUAL SECTIONS ---
    if visual_data:
        # 0. Portfolio Snapshot
        if 'portfolio_snapshot' in visual_data and hasattr(visual_data['portfolio_snapshot'], 'columns'):
            df = visual_data['portfolio_snapshot']
            
            # Filter Logic
            cols = []
            for c in df.columns:
                 c_norm = str(c).replace('\n', ' ').replace('  ', ' ').strip()
                 if "In Place Eff. Rate Prior Month" not in c_norm:
                     cols.append(c)
            
            # Groups matching UI (5 Groups)
            groups = [
                (cols[:5], "Property Details"),
                (cols[5:10], "Cur. Mnth. Operations - Financial Based"),
                (cols[10:16], "NOI - % Variance"),
                (cols[16:22], "Revenue - % Variance"),
                (cols[22:], "Expenses - % Variance")
            ]
            
            for g_cols, title in groups:
                if not g_cols: continue
                
                doc.add_heading(title, level=2)
                sub_df = df[g_cols]
                
                table = doc.add_table(rows=len(sub_df)+1, cols=len(sub_df.columns))
                table.style = 'Table Grid'
                table.autofit = False 
                
                # Headers
                for i, col in enumerate(sub_df.columns):
                    cell = table.rows[0].cells[i]
                    cell.text = str(col)
                    # Bold Header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(8) # Smaller font for headers
                            
                    # Set width (heuristic)
                    # table.columns[i].width = Inches(1.0) # Optional
                    
                # Rows
                for r_idx, row in enumerate(sub_df.itertuples(index=False)):
                    for c_idx, val in enumerate(row):
                        cell = table.rows[r_idx+1].cells[c_idx]
                        cell.text = str(val)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8) # Smaller content font
                
                doc.add_paragraph() # Spacer

        # 1. KPI Snapshot
        if 'kpi_data' in visual_data and hasattr(visual_data['kpi_data'], 'columns'):
            doc.add_heading('KPI SNAPSHOT', level=1)
            df = visual_data['kpi_data']
            table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
            table.style = 'Table Grid'
            
            # Headers
            for i, col in enumerate(df.columns):
                cell = table.rows[0].cells[i]
                cell.text = str(col)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
                
            # Rows
            for r_idx, row in enumerate(df.itertuples(index=False)):
                for c_idx, val in enumerate(row):
                    cell = table.rows[r_idx+1].cells[c_idx]
                    cell.text = str(val)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                    
        # 2. Financial Data
        if 'financial_data' in visual_data and hasattr(visual_data['financial_data'], 'columns'):
            doc.add_heading('MONTHLY FINANCIAL DATA', level=1)
            df = visual_data['financial_data']
            table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
            table.style = 'Table Grid'
            
            # Headers
            for i, col in enumerate(df.columns):
                cell = table.rows[0].cells[i]
                cell.text = str(col)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
            
            # Rows
            for r_idx, row in enumerate(df.itertuples(index=False)):
                for c_idx, val in enumerate(row):
                    cell = table.rows[r_idx+1].cells[c_idx]
                    if isinstance(val, (int, float)):
                        cell.text = f"{val:,.0f}"
                    else:
                        cell.text = str(val)
                    
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)

    # --- AI ANALYSIS ---
    doc.add_heading('AI ANALYSIS & RECOMMENDATIONS', level=1)
    
    # Budget Variances
    if 'budget_variances' in processed_output:
        doc.add_heading('Budget Variances', level=2)
        bv = processed_output['budget_variances']
        for cat, items in bv.items():
            if not items: continue
            doc.add_heading(cat, level=3)
            for item in items:
                p = doc.add_paragraph()
                p.add_run(f"{item.get('metric')}").bold = True
                p.add_run(f"\nActual: ${item.get('actual', 0):,} | Budget: ${item.get('budget', 0):,} | Var: {item.get('variance_pct', 0)}%")
                for q in item.get('questions', []):
                    doc.add_paragraph(f"• {q}", style='List Bullet')
                    
    # Trailing Anomalies
    if 'trailing_anomalies' in processed_output:
        doc.add_heading('Trailing Anomalies', level=2)
        ta = processed_output['trailing_anomalies']
        for cat, items in ta.items():
            if not items: continue
            doc.add_heading(cat, level=3)
            for item in items:
                p = doc.add_paragraph()
                p.add_run(f"{item.get('metric')}").bold = True
                p.add_run(f"\nCurrent: ${item.get('current', 0):,} | T3 Avg: ${item.get('t3_avg', 0):,} | Dev: {item.get('deviation_pct', 0)}%")
                for q in item.get('questions', []):
                    doc.add_paragraph(f"• {q}", style='List Bullet')

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_enhanced_report(processed_output):
    """Generate an enhanced formatted report from processed output"""
    analysis = processed_output["analysis"]
    metadata = processed_output["metadata"]
    property_info = processed_output["property_info"]
    quality = processed_output["quality_metrics"]
    
    timestamp = datetime.fromisoformat(processed_output["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")
    
    report = f"""
T12 PROPERTY ANALYSIS REPORT
============================

PROPERTY INFORMATION
====================
Property Name: {property_info.get('name', 'N/A')}
Property Address: {property_info.get('address', 'N/A')}
Report Generated: {timestamp}

ANALYSIS SUMMARY
================
Quality Score: {quality['overall_score']}/100 ({quality['quality_level']})
Strategic Questions: {len(analysis['strategic_questions'])}
Recommendations: {len(analysis['recommendations'])}
Concerning Trends: {len(analysis['concerning_trends'])}

STRATEGIC MANAGEMENT QUESTIONS
==============================
"""
    
    for i, question in enumerate(analysis["strategic_questions"], 1):
        report += f"{i}. {question}\n"
    
    report += f"""
ACTIONABLE RECOMMENDATIONS
==========================
"""
    
    for i, rec in enumerate(analysis["recommendations"], 1):
        report += f"{i}. {rec}\n"
    
    report += f"""
CONCERNING TRENDS
=================
"""
    
    for i, concern in enumerate(analysis["concerning_trends"], 1):
        report += f"{i}. {concern}\n"
    
    if quality["recommendations"]:
        report += f"""
QUALITY IMPROVEMENT SUGGESTIONS
===============================
"""
        for rec in quality["recommendations"]:
            report += f"• {rec}\n"
    
    report += f"""

TECHNICAL DETAILS
=================
Response Length: {metadata['response_length']} characters
Processing Timestamp: {metadata['processing_timestamp']}

---
Generated by AI-Driven T12 Analysis Tool
Quality Control System v1.0
"""
    
    return report

def generate_text_report(processed_output):
    """Wrapper to generate text report (bytes) from enhanced report string."""
    text_content = generate_enhanced_report(processed_output)
    return text_content.encode('utf-8')

def generate_report(property_name, property_address, kpi_summary, ai_analysis):
    """Generate a simple formatted report for download (legacy function)"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    report = f"""
T12 PROPERTY ANALYSIS REPORT
============================

Property: {property_name}
Address: {property_address}
Generated: {timestamp}

KPI SUMMARY
===========
{kpi_summary}

AI ANALYSIS
===========
{ai_analysis}

---
Generated by AI-Driven T12 Analysis Tool
"""
    return report

def generate_html_download(processed_output, visual_data=None):
    """Generates a full, print-optimized HTML report that mimics the UI exactly."""
    
    # Initialize Generator
    gen = ReportGenerator()
    property_info = processed_output.get("property_info", {})
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    
    # 1. Regenerate Component HTMLs
    # KPI
    if 'kpi_data' in visual_data and 'ytd_kpi_data' in visual_data:
        # We need dictionary format for the generator method
        # But visual_data['kpi_data'] is a DataFrame.
        # We can try to reconstruct or just format the DF directly.
        # EASIER: Re-use the DataFrame to HTML logic? No, ReportGenerator expects dicts for KPI.
        # Let's write a simple DF-to-HTML wrapper that uses the CSS classes.
        kpi_df = visual_data['kpi_data']
        # Use Pandas HTML with class
        kpi_html = kpi_df.to_html(classes='report-table', index=False, escape=False)
    else:
        kpi_html = "<p>No KPI Data</p>"

    # Portfolio Snapshot
    snap_html = "<p>No Snapshot Data</p>"
    if 'portfolio_html' in visual_data:
         # Use pre-generated exact match
         snap_html = visual_data['portfolio_html']
    elif 'portfolio_snapshot' in visual_data:
        # Fallback to generator (requires property name workaround if called)
        try:
             snap_html = gen.generate_portfolio_table(visual_data['portfolio_snapshot'])
        except:
             snap_html = "<p>Error extracting snapshot</p>"
        
    # Financials
    if 'financial_data' in visual_data:
        fin_html = gen.generate_financial_table(visual_data['financial_data'])
    else:
        fin_html = "<p>No Financial Data</p>"
        
    # AI Analysis (Text)
    ai_html = ""
    # Budget Variances
    if 'budget_variances' in processed_output:
        ai_html += "<h3>Budget Variances</h3>"
        for cat, items in processed_output['budget_variances'].items():
            if items:
                ai_html += f"<h4>{cat}</h4><ul>"
                for item in items:
                    ai_html += f"<li><b>{item.get('metric')}</b>: Var {item.get('variance_pct')}% (Act ${item.get('actual',0):,} vs Bud ${item.get('budget',0):,})</li>"
                ai_html += "</ul>"
                
    # Trailing Anomalies
    if 'trailing_anomalies' in processed_output:
        ai_html += "<h3>Trailing Anomalies</h3>"
        for cat, items in processed_output['trailing_anomalies'].items():
            if items:
                ai_html += f"<h4>{cat}</h4><ul>"
                for item in items:
                    ai_html += f"<li><b>{item.get('metric')}</b>: Dev {item.get('deviation_pct')}% (Cur ${item.get('current',0):,} vs T3 ${item.get('t3_avg',0):,})</li>"
                ai_html += "</ul>"

    # Assemble Full Page
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>T12 Analysis - {property_info.get('name', 'Property')}</title>
        {gen.css_styles}
        <style>
            body {{
                background-color: #262730;
                color: #ffffff;
                font-family: 'Segoe UI', sans-serif;
                padding: 40px;
            }}
            .container {{
                max_width: 1200px;
                margin: 0 auto;
            }}
            h1, h2, h3, h4 {{ color: #ffffff; }}
            .header-info {{
                margin-bottom: 30px;
                border-bottom: 1px solid #444;
                padding-bottom: 20px;
            }}
            /* PRINT OPTIMIZATIONS */
            @media print {{
                body {{
                    background-color: #ffffff !important; /* Optional: Print white or keep dark? User liked "Screen" look */
                    /* Actually user said "Actual Page on Screen" which is dark. 
                       But browsers usually force white BG for print to save ink.
                       Force background graphics */
                    -webkit-print-color-adjust: exact;
                    print-color-adjust: exact; 
                }}
                .report-table {{
                    box-shadow: none;
                    break-inside: avoid;
                }}
                /* Force Dark Mode Print if requested */
                body {{
                    background-color: #262730 !important;
                    color: #ffffff !important;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header-info">
                <h1>T12 PROPERTY ANALYSIS REPORT</h1>
                <h2>{property_info.get('name', 'N/A')}</h2>
                <p><strong>Report Period:</strong> {property_info.get('report_period', 'N/A')}</p>
                <p>Generated: {timestamp}</p>
            </div>
            
            <h2>PORTFOLIO SNAPSHOT</h2>
            {snap_html}
            
            <h2>MONTHLY FINANCIAL DATA</h2>
            {fin_html}
            
            <h2>AI ANALYSIS & RECOMMENDATIONS</h2>
            {ai_html}
            
            <p style="margin-top: 50px; font-size: 0.8em; color: #888;">Generated by T12 Analysis Tool</p>
        </div>
        <script>
            // Optional: Auto-print
            // window.print();
        </script>
    </body>
    </html>
    """
    
    return full_html.encode('utf-8')
