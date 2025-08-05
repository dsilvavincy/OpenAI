"""
Report generation utilities for PDF, Word, and text exports
"""
import io
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

def generate_pdf_report(processed_output):
    """Generate a professional PDF report from processed output"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=20,
        textColor=colors.darkblue
    )
    
    # Build content
    story = []
    
    # Title
    story.append(Paragraph("T12 PROPERTY ANALYSIS REPORT", title_style))
    story.append(Spacer(1, 20))
    
    # Property Information
    property_info = processed_output["property_info"]
    story.append(Paragraph("PROPERTY INFORMATION", heading_style))
    property_data = [
        ["Property Name:", property_info.get('name', 'N/A')],
        ["Property Address:", property_info.get('address', 'N/A')],
        ["Report Generated:", datetime.fromisoformat(processed_output["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")]
    ]
    property_table = Table(property_data, colWidths=[2*inch, 4*inch])
    property_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
    ]))
    story.append(property_table)
    story.append(Spacer(1, 20))
    
    # Analysis Summary
    quality = processed_output["quality_metrics"]
    analysis = processed_output["analysis"]
    
    story.append(Paragraph("ANALYSIS SUMMARY", heading_style))
    summary_data = [
        ["Quality Score:", f"{quality['overall_score']}/100 ({quality['quality_level']})"],
        ["Strategic Questions:", str(len(analysis['strategic_questions']))],
        ["Recommendations:", str(len(analysis['recommendations']))],
        ["Concerning Trends:", str(len(analysis['concerning_trends']))]
    ]
    summary_table = Table(summary_data, colWidths=[2*inch, 4*inch])
    summary_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 20))
    
    # Strategic Questions
    story.append(Paragraph("STRATEGIC MANAGEMENT QUESTIONS", heading_style))
    for i, question in enumerate(analysis["strategic_questions"], 1):
        story.append(Paragraph(f"{i}. {question}", styles['Normal']))
        story.append(Spacer(1, 10))
    
    story.append(Spacer(1, 20))
    
    # Recommendations
    story.append(Paragraph("ACTIONABLE RECOMMENDATIONS", heading_style))
    for i, rec in enumerate(analysis["recommendations"], 1):
        story.append(Paragraph(f"{i}. {rec}", styles['Normal']))
        story.append(Spacer(1, 10))
    
    story.append(Spacer(1, 20))
    
    # Concerning Trends
    story.append(Paragraph("CONCERNING TRENDS", heading_style))
    for i, concern in enumerate(analysis["concerning_trends"], 1):
        story.append(Paragraph(f"{i}. {concern}", styles['Normal']))
        story.append(Spacer(1, 10))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_word_report(processed_output):
    """Generate a professional Word document report from processed output"""
    doc = Document()
    
    # Title
    title = doc.add_heading('T12 PROPERTY ANALYSIS REPORT', 0)
    title.alignment = 1  # Center alignment
    
    # Property Information
    doc.add_heading('PROPERTY INFORMATION', level=1)
    property_info = processed_output["property_info"]
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells
    cells[0].text = 'Property Name:'
    cells[1].text = property_info.get('name', 'N/A')
    
    cells = table.rows[1].cells
    cells[0].text = 'Property Address:'
    cells[1].text = property_info.get('address', 'N/A')
    
    cells = table.rows[2].cells
    cells[0].text = 'Report Generated:'
    cells[1].text = datetime.fromisoformat(processed_output["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")
    
    doc.add_paragraph()
    
    # Analysis Summary
    doc.add_heading('ANALYSIS SUMMARY', level=1)
    quality = processed_output["quality_metrics"]
    analysis = processed_output["analysis"]
    
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_data = [
        ['Quality Score:', f"{quality['overall_score']}/100 ({quality['quality_level']})"],
        ['Strategic Questions:', str(len(analysis['strategic_questions']))],
        ['Recommendations:', str(len(analysis['recommendations']))],
        ['Concerning Trends:', str(len(analysis['concerning_trends']))]
    ]
    
    for i, (key, value) in enumerate(summary_data):
        cells = summary_table.rows[i].cells
        cells[0].text = key
        cells[1].text = value
    
    doc.add_paragraph()
    
    # Strategic Questions
    doc.add_heading('STRATEGIC MANAGEMENT QUESTIONS', level=1)
    for i, question in enumerate(analysis["strategic_questions"], 1):
        doc.add_paragraph(f"{i}. {question}")
    
    # Recommendations
    doc.add_heading('ACTIONABLE RECOMMENDATIONS', level=1)
    for i, rec in enumerate(analysis["recommendations"], 1):
        doc.add_paragraph(f"{i}. {rec}")
    
    # Concerning Trends
    doc.add_heading('CONCERNING TRENDS', level=1)
    for i, concern in enumerate(analysis["concerning_trends"], 1):
        doc.add_paragraph(f"{i}. {concern}")
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_enhanced_report(processed_output):
    """Generate an enhanced formatted report from processed output"""
    analysis = processed_output["analysis"]
    metadata = processed_output["metadata"]
    property_info = processed_output["property_info"]
    quality = processed_output["quality_metrics"]
    
    timestamp = datetime.fromisoformat(processed_output["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")
    
    report = f"""
T12 PROPERTY ANALYSIS REPORT
============================

PROPERTY INFORMATION
====================
Property Name: {property_info.get('name', 'N/A')}
Property Address: {property_info.get('address', 'N/A')}
Report Generated: {timestamp}

ANALYSIS SUMMARY
================
Quality Score: {quality['overall_score']}/100 ({quality['quality_level']})
Strategic Questions: {len(analysis['strategic_questions'])}
Recommendations: {len(analysis['recommendations'])}
Concerning Trends: {len(analysis['concerning_trends'])}

STRATEGIC MANAGEMENT QUESTIONS
==============================
"""
    
    for i, question in enumerate(analysis["strategic_questions"], 1):
        report += f"{i}. {question}\n"
    
    report += f"""
ACTIONABLE RECOMMENDATIONS
==========================
"""
    
    for i, rec in enumerate(analysis["recommendations"], 1):
        report += f"{i}. {rec}\n"
    
    report += f"""
CONCERNING TRENDS
=================
"""
    
    for i, concern in enumerate(analysis["concerning_trends"], 1):
        report += f"{i}. {concern}\n"
    
    if quality["recommendations"]:
        report += f"""
QUALITY IMPROVEMENT SUGGESTIONS
===============================
"""
        for rec in quality["recommendations"]:
            report += f"• {rec}\n"
    
    report += f"""

TECHNICAL DETAILS
=================
Response Length: {metadata['response_length']} characters
Processing Timestamp: {metadata['processing_timestamp']}

---
Generated by AI-Driven T12 Analysis Tool
Quality Control System v1.0
"""
    
    return report

def generate_report(property_name, property_address, kpi_summary, ai_analysis):
    """Generate a simple formatted report for download (legacy function)"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    report = f"""
T12 PROPERTY ANALYSIS REPORT
============================

Property: {property_name}
Address: {property_address}
Generated: {timestamp}

KPI SUMMARY
===========
{kpi_summary}

AI ANALYSIS
===========
{ai_analysis}

---
Generated by AI-Driven T12 Analysis Tool
"""
    return report
